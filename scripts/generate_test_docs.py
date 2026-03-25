"""
Script de génération de documents de test pour AtlasMVP.
Génère des documents Word (.docx) réalistes avec données fictives françaises,
organisés par catégorie, et convertit certains en PDF.
"""

import os
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert

BASE_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "documents")

CATEGORIES = [
    "contrats", "baux", "factures", "rapports_financiers",
    "courriers_administratifs", "proces_verbaux", "bons_de_commande",
    "bulletins_de_salaire", "releves_bancaires", "conventions",
]

def setup_dirs():
    for cat in CATEGORIES:
        os.makedirs(os.path.join(BASE_DIR, cat), exist_ok=True)

# ────────────────────────────────────────────────
# DONNÉES FICTIVES
# ────────────────────────────────────────────────

ENTREPRISES = [
    {"nom": "Nexoria Solutions SAS", "forme": "SAS", "capital": "150 000", "siren": "512 348 901",
     "siret": "512 348 901 00023", "adresse": "14 rue du Faubourg Saint-Antoine, 75011 Paris",
     "ville": "Paris", "tva_intra": "FR 45 512348901", "iban": "FR76 3000 6000 0112 3456 7890 189",
     "bic": "AGRIFRPP", "banque": "Crédit Agricole", "tel": "01 42 58 33 10", "email": "contact@nexoria-solutions.fr"},
    {"nom": "Altimum Consulting SARL", "forme": "SARL", "capital": "80 000", "siren": "439 812 567",
     "siret": "439 812 567 00041", "adresse": "7 avenue des Gobelins, 75013 Paris",
     "ville": "Paris", "tva_intra": "FR 22 439812567", "iban": "FR76 1027 8022 0000 0203 8460 158",
     "bic": "CMCIFRPP", "banque": "CIC", "tel": "01 53 79 44 20", "email": "direction@altimum-consulting.fr"},
    {"nom": "Groupe Lumière & Associés SA", "forme": "SA", "capital": "500 000", "siren": "378 524 109",
     "siret": "378 524 109 00058", "adresse": "42 boulevard Haussmann, 75009 Paris",
     "ville": "Paris", "tva_intra": "FR 18 378524109", "iban": "FR76 3000 3032 0000 0503 2145 862",
     "bic": "SOGEFRPP", "banque": "Société Générale", "tel": "01 45 22 88 00", "email": "info@groupe-lumiere.fr"},
    {"nom": "Terralpha SAS", "forme": "SAS", "capital": "200 000", "siren": "621 047 335",
     "siret": "621 047 335 00012", "adresse": "3 rue Nationale, 69001 Lyon",
     "ville": "Lyon", "tva_intra": "FR 37 621047335", "iban": "FR76 1780 6004 0000 1234 5678 910",
     "bic": "BNPAFRPP", "banque": "BNP Paribas", "tel": "04 72 41 33 50", "email": "contact@terralpha.fr"},
    {"nom": "Oryon Technologies SA", "forme": "SA", "capital": "1 000 000", "siren": "552 148 774",
     "siret": "552 148 774 00034", "adresse": "25 allée de la Défense, 92400 Courbevoie",
     "ville": "Courbevoie", "tva_intra": "FR 69 552148774", "iban": "FR76 3000 4000 0300 0503 8920 147",
     "bic": "BNPAFRPP", "banque": "BNP Paribas", "tel": "01 49 01 55 55", "email": "contact@oryon-tech.fr"},
    {"nom": "Novéa Services SARL", "forme": "SARL", "capital": "30 000", "siren": "801 223 456",
     "siret": "801 223 456 00018", "adresse": "18 rue de la Paix, 13001 Marseille",
     "ville": "Marseille", "tva_intra": "FR 10 801223456", "iban": "FR76 1558 9200 0109 2345 6789 012",
     "bic": "CEPAFR21", "banque": "Caisse d'Épargne", "tel": "04 91 44 12 80", "email": "info@novea-services.fr"},
    {"nom": "Deltaform Groupe SAS", "forme": "SAS", "capital": "350 000", "siren": "487 659 201",
     "siret": "487 659 201 00027", "adresse": "9 rue du Commerce, 31000 Toulouse",
     "ville": "Toulouse", "tva_intra": "FR 55 487659201", "iban": "FR76 1660 6001 0080 0500 1234 567",
     "bic": "CMCIFRPP", "banque": "CIC Toulouse", "tel": "05 61 23 44 70", "email": "direction@deltaform.fr"},
    {"nom": "Prismédia Communication SA", "forme": "SA", "capital": "250 000", "siren": "339 887 542",
     "siret": "339 887 542 00065", "adresse": "56 cours Pasteur, 33000 Bordeaux",
     "ville": "Bordeaux", "tva_intra": "FR 44 339887542", "iban": "FR76 1025 7020 0000 3456 7890 123",
     "bic": "LCFRFRP1", "banque": "LCL", "tel": "05 56 33 22 11", "email": "contact@prismedia.fr"},
    {"nom": "Veridis Audit & Finance SARL", "forme": "SARL", "capital": "60 000", "siren": "710 345 987",
     "siret": "710 345 987 00039", "adresse": "2 place du Général de Gaulle, 67000 Strasbourg",
     "ville": "Strasbourg", "tva_intra": "FR 88 710345987", "iban": "FR76 1466 7000 0100 1234 5678 901",
     "bic": "CEPAFR21", "banque": "Caisse d'Épargne Alsace", "tel": "03 88 22 55 44", "email": "audit@veridis.fr"},
    {"nom": "Apexline Industrie SAS", "forme": "SAS", "capital": "750 000", "siren": "443 910 268",
     "siret": "443 910 268 00051", "adresse": "Zone Industrielle Nord, 59000 Lille",
     "ville": "Lille", "tva_intra": "FR 77 443910268", "iban": "FR76 3000 3000 0300 0403 8910 257",
     "bic": "SOGEFRPP", "banque": "Société Générale Nord", "tel": "03 20 44 33 22", "email": "info@apexline.fr"},
]

PERSONNES = [
    {"nom": "Martin Dupont", "prenom": "Martin", "civilite": "M.", "poste": "Directeur Général",
     "email": "m.dupont@nexoria-solutions.fr", "tel": "06 12 34 56 78"},
    {"nom": "Sophie Laurent", "prenom": "Sophie", "civilite": "Mme", "poste": "Directrice Financière",
     "email": "s.laurent@altimum-consulting.fr", "tel": "06 23 45 67 89"},
    {"nom": "Karim Benali", "prenom": "Karim", "civilite": "M.", "poste": "Directeur des Opérations",
     "email": "k.benali@groupe-lumiere.fr", "tel": "06 34 56 78 90"},
    {"nom": "Isabelle Moreau", "prenom": "Isabelle", "civilite": "Mme", "poste": "Responsable Juridique",
     "email": "i.moreau@terralpha.fr", "tel": "06 45 67 89 01"},
    {"nom": "Thomas Renard", "prenom": "Thomas", "civilite": "M.", "poste": "Président",
     "email": "t.renard@oryon-tech.fr", "tel": "06 56 78 90 12"},
    {"nom": "Nathalie Petit", "prenom": "Nathalie", "civilite": "Mme", "poste": "Directrice Commerciale",
     "email": "n.petit@novea-services.fr", "tel": "06 67 89 01 23"},
    {"nom": "Alexandre Mercier", "prenom": "Alexandre", "civilite": "M.", "poste": "DRH",
     "email": "a.mercier@deltaform.fr", "tel": "06 78 90 12 34"},
    {"nom": "Céline Fontaine", "prenom": "Céline", "civilite": "Mme", "poste": "Directrice Générale",
     "email": "c.fontaine@prismedia.fr", "tel": "06 89 01 23 45"},
    {"nom": "Pierre Girard", "prenom": "Pierre", "civilite": "M.", "poste": "Directeur Technique",
     "email": "p.girard@veridis.fr", "tel": "06 90 12 34 56"},
    {"nom": "Amandine Leroy", "prenom": "Amandine", "civilite": "Mme", "poste": "Présidente",
     "email": "a.leroy@apexline.fr", "tel": "06 01 23 45 67"},
]

VILLES = ["Paris", "Lyon", "Marseille", "Toulouse", "Bordeaux", "Lille", "Strasbourg", "Nantes", "Nice", "Rennes"]
TRIBUNAUX = ["Paris", "Lyon", "Marseille", "Bordeaux", "Lille"]

def date_str(offset_days=0):
    d = datetime.now() + timedelta(days=offset_days)
    return d.strftime("%d/%m/%Y")

def add_header(doc, org_name, doc_type, ref):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(org_name)
    run.bold = True
    run.font.size = Pt(16)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(doc_type)
    r2.bold = True
    r2.font.size = Pt(13)

    doc.add_paragraph(f"Référence : {ref}")
    doc.add_paragraph(f"Date : {date_str()}")
    doc.add_paragraph()

def add_signature_block(doc, part_a, rep_a, qual_a, part_b, rep_b, qual_b, ville):
    doc.add_paragraph()
    doc.add_paragraph(f"Fait à {ville}, le {date_str()}")
    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    c1, c2 = table.rows[0].cells
    c1.text = f"Pour {part_a}\nSignature : _______________\nNom : {rep_a}\nFonction : {qual_a}"
    c2.text = f"Pour {part_b}\nSignature : _______________\nNom : {rep_b}\nFonction : {qual_b}"

# ────────────────────────────────────────────────
# CONTRATS DE PRESTATION (10 docs)
# ────────────────────────────────────────────────

CONTRAT_PRESTATIONS = [
    ("Développement d'une plateforme e-commerce", "CONT-PREST-2024-001", "4 mois", 48000),
    ("Audit de sécurité informatique", "CONT-PREST-2024-002", "6 semaines", 12500),
    ("Mission de conseil en stratégie digitale", "CONT-PREST-2024-003", "3 mois", 30000),
    ("Programme de formation professionnelle", "CONT-PREST-2024-004", "2 mois", 8500),
    ("Maintenance applicative et support", "CONT-PREST-2024-005", "12 mois", 60000),
    ("Conception graphique – Refonte identité visuelle", "CONT-PREST-2024-006", "6 semaines", 9200),
    ("Traduction et localisation – 3 langues", "CONT-PREST-2024-007", "3 semaines", 4800),
    ("Analyse de données et tableaux de bord BI", "CONT-PREST-2024-008", "2 mois", 22000),
    ("Rédaction technique – Documentation produit", "CONT-PREST-2024-009", "1 mois", 6000),
    ("Intégration ERP – Module comptabilité", "CONT-PREST-2024-010", "5 mois", 75000),
]

def gen_contrat_prestation(idx, objet, ref, duree, ht):
    client = ENTREPRISES[idx % len(ENTREPRISES)]
    prest = ENTREPRISES[(idx + 3) % len(ENTREPRISES)]
    rep_client = PERSONNES[idx % len(PERSONNES)]
    rep_prest = PERSONNES[(idx + 2) % len(PERSONNES)]
    tva = round(ht * 0.20, 2)
    ttc = round(ht + tva, 2)
    acompte = round(ht * 0.30, 2)

    doc = Document()
    add_header(doc, client["nom"], f"CONTRAT DE PRESTATION DE SERVICES\n{objet}", ref)

    doc.add_heading("ENTRE LES SOUSSIGNÉS", level=2)
    doc.add_paragraph(
        f"{client['nom']}, {client['forme']} au capital de {client['capital']} €, "
        f"immatriculée sous le numéro SIREN {client['siren']}, dont le siège social est situé au "
        f"{client['adresse']}, représentée par {rep_client['nom']}, "
        f"agissant en qualité de {rep_client['poste']}, ci-après désignée « le CLIENT »,"
    )
    doc.add_paragraph("ET")
    doc.add_paragraph(
        f"{prest['nom']}, {prest['forme']} au capital de {prest['capital']} €, "
        f"immatriculée sous le numéro SIREN {prest['siren']}, dont le siège social est situé au "
        f"{prest['adresse']}, représentée par {rep_prest['nom']}, "
        f"agissant en qualité de {rep_prest['poste']}, ci-après désignée « le PRESTATAIRE »,"
    )

    doc.add_heading("ARTICLE 1 – OBJET", level=2)
    doc.add_paragraph(
        f"Le présent contrat a pour objet la réalisation de prestations de services portant sur : "
        f"{objet}, telles que décrites en Annexe A du présent contrat."
    )

    doc.add_heading("ARTICLE 2 – DURÉE", level=2)
    doc.add_paragraph(
        f"Le contrat prend effet à compter du {date_str()} pour une durée de {duree}. "
        f"Il prend fin automatiquement à l'échéance, sauf reconduction expresse par écrit des deux parties."
    )

    doc.add_heading("ARTICLE 3 – RÉMUNÉRATION", level=2)
    doc.add_paragraph(
        f"En contrepartie des prestations réalisées, le CLIENT versera au PRESTATAIRE "
        f"la somme forfaitaire de {ht:,.2f} € HT, soit {ttc:,.2f} € TTC (TVA à 20 %)."
    )
    doc.add_paragraph("Les modalités de paiement sont les suivantes :")
    doc.add_paragraph(f"  – 30 % à la signature du contrat : {acompte:,.2f} € HT", style="List Bullet")
    doc.add_paragraph(f"  – 70 % à la livraison finale validée par le CLIENT : {ht - acompte:,.2f} € HT.", style="List Bullet")
    doc.add_paragraph("Paiement par virement bancaire à 30 jours date de facture.")

    doc.add_heading("ARTICLE 4 – CONFIDENTIALITÉ", level=2)
    doc.add_paragraph(
        "Chaque partie s'engage à garder confidentielles toutes informations reçues de l'autre "
        "partie dans le cadre du présent contrat, pour une durée de 5 ans après son expiration."
    )

    doc.add_heading("ARTICLE 5 – PROPRIÉTÉ INTELLECTUELLE", level=2)
    doc.add_paragraph(
        f"Les livrables produits dans le cadre du présent contrat seront la propriété exclusive "
        f"du CLIENT ({client['nom']}) après paiement intégral des sommes dues."
    )

    doc.add_heading("ARTICLE 6 – RÉSILIATION", level=2)
    doc.add_paragraph(
        "Chaque partie peut résilier le présent contrat en cas de manquement grave de l'autre "
        "partie à ses obligations, après mise en demeure restée sans effet pendant 15 jours ouvrés."
    )

    doc.add_heading("ARTICLE 7 – LOI APPLICABLE ET JURIDICTION", level=2)
    tribunal = TRIBUNAUX[idx % len(TRIBUNAUX)]
    doc.add_paragraph(
        f"Le présent contrat est soumis au droit français. En cas de litige, les parties conviennent "
        f"de tenter une médiation amiable avant tout recours judiciaire. À défaut, le Tribunal de "
        f"Commerce de {tribunal} sera seul compétent."
    )

    add_signature_block(doc, client["nom"], rep_client["nom"], rep_client["poste"],
                        prest["nom"], rep_prest["nom"], rep_prest["poste"], client["ville"])
    path = os.path.join(BASE_DIR, "contrats", f"contrat_prestation_{idx:03d}.docx")
    doc.save(path)
    return path

# ────────────────────────────────────────────────
# CONTRATS DE BAIL (10 docs)
# ────────────────────────────────────────────────

BAUX_DATA = [
    ("BAIL-COM-2024-001", "Bail commercial – Local 75 m²", "9 ans", 1800, "75 m²",
     "8 rue du Temple, 75004 Paris", "Local commercial en rez-de-chaussée avec vitrine sur rue", 5400),
    ("BAIL-RES-2024-001", "Bail d'habitation – Appartement T3 68 m²", "3 ans", 950, "68 m²",
     "12 avenue Victor Hugo, 75016 Paris", "Appartement T3 au 3ème étage avec ascenseur", 1900),
    ("BAIL-PRO-2024-001", "Bail professionnel – Bureau 45 m²", "6 ans", 1200, "45 m²",
     "23 rue de Rivoli, 75001 Paris", "Bureau fermé au 2ème étage, charge commune incluse", 2400),
    ("BAIL-COM-2024-002", "Bail commercial – Entrepôt 500 m²", "9 ans", 3500, "500 m²",
     "Zone Logistique Nord, 93400 Saint-Ouen", "Entrepôt de stockage avec quai de chargement", 7000),
    ("BAIL-RES-2024-002", "Bail meublé – Studio 28 m²", "1 an renouvelable", 820, "28 m²",
     "5 impasse des Lilas, 75020 Paris", "Studio meublé au 1er étage, double vitrage", 1640),
    ("BAIL-PREC-2024-001", "Bail dérogatoire – Local commercial 40 m²", "23 mois", 900, "40 m²",
     "44 boulevard Beaumarchais, 75011 Paris", "Pop-up store en rez-de-chaussée, accès PMR", 1800),
    ("BAIL-COM-2024-003", "Bail commercial – Restaurant 120 m²", "9 ans", 2900, "120 m²",
     "18 place du Capitole, 31000 Toulouse", "Local restaurant avec terrasse de 30 m², hotte installée", 8700),
    ("BAIL-RES-2024-003", "Bail d'habitation – Maison T4 110 m²", "3 ans", 1450, "110 m²",
     "7 allée des Chênes, 33600 Pessac", "Maison individuelle avec jardin 200 m², garage", 2900),
    ("BAIL-PRO-2024-002", "Bail professionnel – Cabinet médical 60 m²", "6 ans", 1800, "60 m²",
     "15 rue de la Santé, 67000 Strasbourg", "Cabinet médical aménagé, salle d'attente et 2 salles de soins", 3600),
    ("BAIL-COM-2024-004", "Bail commercial – Boutique 60 m²", "9 ans", 2100, "60 m²",
     "32 rue de la République, 69001 Lyon", "Boutique avec réserve de 15 m² et vitrine double", 6300),
]

def gen_bail(idx, ref, titre, duree, loyer, surface, adresse_bien, desc, depot):
    bailleur = ENTREPRISES[idx % len(ENTREPRISES)]
    preneur = ENTREPRISES[(idx + 5) % len(ENTREPRISES)]
    rep_bailleur = PERSONNES[idx % len(PERSONNES)]
    rep_preneur = PERSONNES[(idx + 4) % len(PERSONNES)]

    doc = Document()
    add_header(doc, bailleur["nom"], titre.upper(), ref)

    doc.add_heading("PARTIES", level=2)
    doc.add_paragraph(
        f"BAILLEUR : {bailleur['nom']}, {bailleur['forme']}, SIREN {bailleur['siren']}, "
        f"dont le siège est au {bailleur['adresse']}, "
        f"représenté(e) par {rep_bailleur['nom']}, {rep_bailleur['poste']},"
    )
    doc.add_paragraph(
        f"PRENEUR : {preneur['nom']}, {preneur['forme']}, SIREN {preneur['siren']}, "
        f"dont le siège est au {preneur['adresse']}, "
        f"représenté(e) par {rep_preneur['nom']}, {rep_preneur['poste']}."
    )

    doc.add_heading("ARTICLE 1 – DÉSIGNATION DES LOCAUX", level=2)
    doc.add_paragraph(
        f"Le BAILLEUR loue au PRENEUR les locaux situés au {adresse_bien}, "
        f"d'une surface de {surface}, comprenant : {desc}."
    )

    doc.add_heading("ARTICLE 2 – DURÉE", level=2)
    doc.add_paragraph(
        f"Le présent bail est consenti pour une durée de {duree}, "
        f"à compter du {date_str()} jusqu'au {date_str(365 * 3)}."
    )

    doc.add_heading("ARTICLE 3 – LOYER ET CHARGES", level=2)
    doc.add_paragraph(
        f"Le loyer mensuel est fixé à {loyer:,} € HT, payable le 1er de chaque mois par virement. "
        f"Les charges locatives récupérables s'élèvent à {round(loyer * 0.12):,} €/mois (provision). "
        f"Dépôt de garantie : {depot:,} €."
    )

    doc.add_heading("ARTICLE 4 – DESTINATION DES LOCAUX", level=2)
    doc.add_paragraph(
        f"Les locaux sont destinés exclusivement à l'exercice de l'activité de {preneur['nom']}. "
        "Toute modification d'usage est soumise à l'accord préalable et écrit du BAILLEUR."
    )

    doc.add_heading("ARTICLE 5 – ÉTAT DES LIEUX", level=2)
    doc.add_paragraph(
        f"Un état des lieux d'entrée contradictoire sera établi le {date_str(7)}. "
        "Un état des lieux de sortie sera réalisé à l'expiration du bail. "
        "Tout état des lieux sera annexé au présent bail."
    )

    doc.add_heading("ARTICLE 6 – TRAVAUX", level=2)
    doc.add_paragraph(
        "Le PRENEUR ne peut effectuer de travaux de transformation sans l'accord écrit préalable du BAILLEUR. "
        "Les travaux d'entretien courant sont à la charge du PRENEUR (cf. décret n°87-712 du 26/08/1987)."
    )

    doc.add_heading("ARTICLE 7 – ASSURANCES", level=2)
    doc.add_paragraph(
        "Le PRENEUR s'engage à souscrire et maintenir, pendant toute la durée du bail, "
        "une assurance multirisque couvrant sa responsabilité civile. "
        f"Justificatif à fournir annuellement au BAILLEUR ({rep_bailleur['email']})."
    )

    add_signature_block(doc, bailleur["nom"], rep_bailleur["nom"], rep_bailleur["poste"],
                        preneur["nom"], rep_preneur["nom"], rep_preneur["poste"], bailleur["ville"])
    path = os.path.join(BASE_DIR, "baux", f"bail_{idx:03d}.docx")
    doc.save(path)
    return path

# ────────────────────────────────────────────────
# FACTURES (10 docs)
# ────────────────────────────────────────────────

FACTURES_DATA = [
    ("FAC-2024-0001", "Développement site web e-commerce – Phase 1", 18500),
    ("FAC-2024-0002", "Formation Cybersécurité – 2 jours – 8 participants", 3200),
    ("FAC-2024-0003", "Licences logicielles annuelles Suite Office x10", 4500),
    ("FAC-2024-0004", "Maintenance infrastructure serveurs – T3 2024", 2800),
    ("FAC-2024-0005", "Mission de conseil RH et recrutement cadres", 5600),
    ("FAC-2024-0006", "Impression et reliure – 500 rapports annuels", 1200),
    ("FAC-2024-0007", "Fournitures de bureau – Commande mensuelle octobre", 480),
    ("FAC-2024-0008", "Audit comptable et fiscal – Exercice 2023", 7500),
    ("FAC-2024-0009", "Location salle de conférence 200 m² – Novembre 2024", 950),
    ("FAC-2024-0010", "Traduction documents – Anglais/Français – 45 000 mots", 2100),
]

def gen_facture(idx, num, objet, ht):
    fournisseur = ENTREPRISES[idx % len(ENTREPRISES)]
    client = ENTREPRISES[(idx + 4) % len(ENTREPRISES)]
    tva = round(ht * 0.20, 2)
    ttc = round(ht + tva, 2)

    doc = Document()
    add_header(doc, fournisseur["nom"], "FACTURE", num)

    table_coords = doc.add_table(rows=1, cols=2)
    table_coords.style = "Table Grid"
    left, right = table_coords.rows[0].cells
    left.text = (
        f"ÉMETTEUR :\n{fournisseur['nom']}\n{fournisseur['adresse']}\n"
        f"SIRET : {fournisseur['siret']}\nN° TVA intra : {fournisseur['tva_intra']}\n"
        f"Tél : {fournisseur['tel']}\nEmail : {fournisseur['email']}"
    )
    right.text = (
        f"CLIENT :\n{client['nom']}\n{client['adresse']}\n"
        f"SIRET : {client['siret']}\nN° TVA intra : {client['tva_intra']}"
    )
    doc.add_paragraph()

    doc.add_paragraph(f"Date d'émission : {date_str()}")
    doc.add_paragraph(f"Date d'échéance : {date_str(30)}")
    doc.add_paragraph("Conditions de paiement : 30 jours net, virement bancaire")
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["Description", "Qté", "Prix unitaire HT", "TVA", "Total HT"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    row = table.add_row().cells
    row[0].text = objet
    row[1].text = "1"
    row[2].text = f"{ht:,.2f} €"
    row[3].text = "20 %"
    row[4].text = f"{ht:,.2f} €"

    doc.add_paragraph()
    doc.add_paragraph(f"Total HT : {ht:,.2f} €")
    doc.add_paragraph(f"TVA (20 %) : {tva:,.2f} €")
    p = doc.add_paragraph(f"TOTAL TTC : {ttc:,.2f} €")
    p.runs[0].bold = True

    doc.add_paragraph()
    doc.add_paragraph(
        f"Règlement par virement bancaire :\n"
        f"IBAN : {fournisseur['iban']}\nBIC : {fournisseur['bic']}\n"
        f"Banque : {fournisseur['banque']}\n"
        f"Référence à indiquer : {num}"
    )
    doc.add_paragraph(
        "En cas de retard de paiement, des pénalités seront appliquées au taux de 3 fois le "
        "taux d'intérêt légal, ainsi qu'une indemnité forfaitaire de recouvrement de 40 €."
    )

    path = os.path.join(BASE_DIR, "factures", f"facture_{idx:03d}.docx")
    doc.save(path)
    return path

# ────────────────────────────────────────────────
# RAPPORTS FINANCIERS (10 docs)
# ────────────────────────────────────────────────

RAPPORTS_FIN_DATA = [
    ("Rapport financier annuel 2023", "RAP-FIN-2023-ANNUEL", "01/01/2023", "31/12/2023",
     12450000, 11820000, 630000, 850000, -220000),
    ("Rapport financier T1 2024", "RAP-FIN-2024-T1", "01/01/2024", "31/03/2024",
     3200000, 2980000, 220000, 3500000, 2800000),
    ("Rapport financier T2 2024", "RAP-FIN-2024-T2", "01/04/2024", "30/06/2024",
     3450000, 3100000, 350000, 3800000, 2950000),
    ("Rapport financier T3 2024", "RAP-FIN-2024-T3", "01/07/2024", "30/09/2024",
     2980000, 2750000, 230000, 3200000, 2600000),
    ("Budget prévisionnel 2025", "RAP-BUDGET-2025", "01/01/2025", "31/12/2025",
     15000000, 13500000, 1500000, 8500000, 6000000),
    ("Rapport d'audit interne 2024", "RAP-AUDIT-2024", "01/01/2024", "31/12/2024",
     13200000, 12100000, 1100000, 9200000, 7500000),
    ("Analyse de rentabilité – Projet ALPHA", "RAP-RENT-ALPHA-2024", "01/03/2024", "30/09/2024",
     2800000, 2100000, 700000, 1200000, 900000),
    ("Tableau de bord financier – Novembre 2024", "RAP-TDB-NOV-2024", "01/11/2024", "30/11/2024",
     1100000, 980000, 120000, 2400000, 1900000),
    ("Rapport de trésorerie – Exercice 2024", "RAP-TRES-2024", "01/01/2024", "31/12/2024",
     14200000, 13000000, 1200000, 3100000, 2500000),
    ("Rapport d'activité consolidé 2023", "RAP-CONSOL-2023", "01/01/2023", "31/12/2023",
     25600000, 23800000, 1800000, 12500000, 9800000),
]

def gen_rapport_financier(idx, titre, ref, debut, fin, ca, charges, ebit, total_actif, capitaux):
    org = ENTREPRISES[idx % len(ENTREPRISES)]
    daf = PERSONNES[(idx + 1) % len(PERSONNES)]
    pdg = PERSONNES[idx % len(PERSONNES)]
    result_net = round(ebit * 0.75, 0)

    doc = Document()
    add_header(doc, org["nom"], titre.upper(), ref)

    doc.add_paragraph(
        f"Préparé par : {daf['nom']}, {daf['poste']}\n"
        f"Approuvé par : {pdg['nom']}, {pdg['poste']}\n"
        f"Période couverte : {debut} au {fin}\n"
        f"Date de diffusion : {date_str()}"
    )
    doc.add_paragraph()

    doc.add_heading("1. SYNTHÈSE EXÉCUTIVE", level=1)
    doc.add_paragraph(
        f"Le présent rapport présente la situation financière de {org['nom']} "
        f"pour la période du {debut} au {fin}."
    )
    doc.add_paragraph(
        f"Chiffre d'affaires : {ca:,} €\n"
        f"Charges d'exploitation : {charges:,} €\n"
        f"Résultat d'exploitation (EBIT) : {ebit:,} €  (marge : {round(ebit/ca*100, 1)} %)\n"
        f"Résultat net estimé : {result_net:,.0f} €\n"
        f"Total actif : {total_actif:,} €\n"
        f"Capitaux propres : {capitaux:,} €"
    )

    doc.add_heading("2. COMPTE DE RÉSULTAT SIMPLIFIÉ", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for i, h in enumerate(["Rubrique", "Montant (€)", "% CA"]):
        table.rows[0].cells[i].text = h
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    rows_cr = [
        ("Produits d'exploitation (CA)", f"{ca:,}", "100,0 %"),
        ("Achats et charges externes", f"-{round(charges*0.55):,}", f"-{round(charges*0.55/ca*100,1)} %"),
        ("Frais de personnel", f"-{round(charges*0.35):,}", f"-{round(charges*0.35/ca*100,1)} %"),
        ("Amortissements et provisions", f"-{round(charges*0.10):,}", f"-{round(charges*0.10/ca*100,1)} %"),
        ("Résultat d'exploitation (EBIT)", f"{ebit:,}", f"{round(ebit/ca*100,1)} %"),
        ("Résultat financier", f"{round(ebit*-0.08):,}", f"{round(ebit*-0.08/ca*100,2)} %"),
        ("Résultat avant impôt", f"{round(ebit*0.92):,}", f"{round(ebit*0.92/ca*100,1)} %"),
        ("Impôt sur les sociétés (25 %)", f"-{round(ebit*0.92*0.25):,}", f"-{round(ebit*0.92*0.25/ca*100,1)} %"),
        ("Résultat net", f"{result_net:,.0f}", f"{round(result_net/ca*100,1)} %"),
    ]
    for row_data in rows_cr:
        r = table.add_row().cells
        for i, v in enumerate(row_data):
            r[i].text = v

    doc.add_heading("3. BILAN SIMPLIFIÉ", level=1)
    immo = round(total_actif * 0.40)
    stocks = round(total_actif * 0.10)
    creances = round(total_actif * 0.25)
    dispo = total_actif - immo - stocks - creances
    dettes_fin = round(total_actif * 0.30)
    dettes_fourn = round(total_actif * 0.15)
    autres_dettes = total_actif - capitaux - dettes_fin - dettes_fourn

    doc.add_paragraph("ACTIF")
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    for label, val in [
        ("Immobilisations nettes", f"{immo:,} €"),
        ("Stocks", f"{stocks:,} €"),
        ("Créances clients", f"{creances:,} €"),
        ("Disponibilités et équivalents", f"{dispo:,} €"),
        ("TOTAL ACTIF", f"{total_actif:,} €"),
    ]:
        r = t.add_row().cells
        r[0].text = label
        r[1].text = val

    doc.add_paragraph()
    doc.add_paragraph("PASSIF")
    t2 = doc.add_table(rows=1, cols=2)
    t2.style = "Table Grid"
    for label, val in [
        ("Capitaux propres", f"{capitaux:,} €"),
        ("Dettes financières", f"{dettes_fin:,} €"),
        ("Dettes fournisseurs", f"{dettes_fourn:,} €"),
        ("Autres dettes", f"{max(autres_dettes,0):,} €"),
        ("TOTAL PASSIF", f"{total_actif:,} €"),
    ]:
        r = t2.add_row().cells
        r[0].text = label
        r[1].text = val

    doc.add_heading("4. FLUX DE TRÉSORERIE", level=1)
    flux_op = round(ebit + charges * 0.10)
    flux_inv = round(-total_actif * 0.05)
    flux_fin = round(-dettes_fin * 0.08)
    doc.add_paragraph(
        f"Flux opérationnels : +{flux_op:,} €\n"
        f"Flux d'investissement : {flux_inv:,} €\n"
        f"Flux de financement : {flux_fin:,} €\n"
        f"Variation nette de trésorerie : {flux_op + flux_inv + flux_fin:,} €"
    )

    doc.add_heading("5. COMMENTAIRES DE DIRECTION", level=1)
    doc.add_paragraph(
        f"La direction de {org['nom']} note une évolution {('favorable' if ebit > 0 else 'défavorable')} "
        f"de la performance opérationnelle sur la période. "
        f"La marge d'exploitation de {round(ebit/ca*100,1)} % est "
        f"{'supérieure' if ebit/ca > 0.05 else 'inférieure'} à l'objectif de 5,0 % fixé en début d'exercice.\n\n"
        f"Les priorités pour la prochaine période incluent : l'optimisation des charges externes, "
        f"l'accélération du recouvrement client et le suivi des investissements en cours."
    )

    path = os.path.join(BASE_DIR, "rapports_financiers", f"rapport_financier_{idx:03d}.docx")
    doc.save(path)
    return path

# ────────────────────────────────────────────────
# COURRIERS ADMINISTRATIFS (10 docs)
# ────────────────────────────────────────────────

COURRIERS_DATA = [
    ("Demande de subvention – Programme Innovation Numérique 2024",
     "COUR-2024-001",
     "Dans le cadre du Programme National d'Innovation Numérique 2024, notre organisation "
     "souhaite soumettre sa candidature pour l'obtention d'une subvention destinée à financer "
     "le développement d'un outil de gestion documentaire basé sur l'intelligence artificielle.",
     "Nous sollicitons une subvention d'un montant de 85 000 €. Le dossier complet de candidature, "
     "incluant le plan de financement et les indicateurs de performance attendus, est joint en annexe.",
     "Nous restons à votre entière disposition pour tout entretien ou complément d'information que vous jugeriez utile."),
    ("Réponse à l'appel d'offres n°AO-2024-0445",
     "COUR-2024-002",
     "Suite à la publication de votre appel d'offres n°AO-2024-0445 relatif à la fourniture de "
     "prestations de conseil en transformation digitale, nous avons l'honneur de vous adresser "
     "notre offre de services.",
     "Notre proposition technique et financière, jointe en annexe, répond point par point au cahier "
     "des charges. Notre équipe de consultants seniors dispose d'une expérience de plus de 10 ans "
     "dans des missions similaires pour des organisations de taille comparable.",
     "Nous espérons que notre candidature retiendra votre attention et vous remercions de l'intérêt que vous portez à notre démarche."),
    ("Réclamation relative à la facture FAC-FOURN-2024-0892",
     "COUR-2024-003",
     "Nous avons bien réceptionné votre facture n°FAC-FOURN-2024-0892 d'un montant de 12 400 € TTC "
     "en date du 15 octobre 2024. Après vérification de celle-ci au regard du bon de commande "
     "BDC-2024-0087, nous avons constaté une anomalie de facturation.",
     "La quantité facturée (50 unités) ne correspond pas à la quantité commandée et livrée (42 unités). "
     "Nous vous demandons en conséquence d'établir un avoir pour la différence de 8 unités, "
     "soit 1 984 € TTC, et de nous transmettre une facture rectificative.",
     "En l'attente de la régularisation, nous suspendons le paiement de cette facture."),
    ("Notification de résiliation du contrat CONT-PREST-2023-018",
     "COUR-2024-004",
     "Par la présente, nous vous informons de notre décision de mettre fin au contrat de prestation "
     "de services n°CONT-PREST-2023-018, conformément aux dispositions de l'article 6 du dit contrat "
     "qui prévoit un préavis de résiliation de 30 jours calendaires.",
     "Cette résiliation prendra effet le 30 novembre 2024. Nous vous remercions de prendre toutes "
     "dispositions nécessaires pour assurer la transmission des livrables en cours et la restitution "
     "de tous les éléments appartenant à notre organisation.",
     "Un décompte des prestations réalisées à la date effective de résiliation vous sera adressé séparément."),
    ("Demande d'agrément – Organisme de formation professionnelle",
     "COUR-2024-005",
     "Dans le cadre du développement de nos activités de formation professionnelle continue, "
     "nous sollicitons auprès de vos services l'agrément d'organisme de formation conformément "
     "aux articles L.6351-1 et suivants du Code du Travail.",
     "Notre dossier de demande d'agrément est joint à ce courrier. Il comprend notre déclaration "
     "d'activité, le programme de nos formations, les qualifications de nos formateurs et "
     "notre bilan pédagogique et financier des 12 derniers mois.",
     "Nous nous tenons disponibles pour vous rencontrer ou vous fournir tout document complémentaire."),
    ("Mise en demeure de règlement – Facture FAC-2024-0234",
     "COUR-2024-006",
     "Malgré nos relances téléphoniques et écrites du 15 et 28 octobre 2024, "
     "nous constatons que la facture n°FAC-2024-0234 d'un montant de 9 750 € TTC, "
     "émise le 1er septembre 2024 et échue le 1er octobre 2024, demeure impayée à ce jour.",
     "Nous vous mettons en demeure de procéder au règlement intégral de cette somme dans un délai "
     "de 8 jours à compter de la réception du présent courrier. À défaut, nous nous verrons "
     "dans l'obligation d'engager une procédure de recouvrement judiciaire, dont les frais "
     "vous seront intégralement facturés.",
     "Des pénalités de retard au taux de 3 fois le taux légal seront également réclamées."),
    ("Demande de prolongation du délai de remise des offres – AO-2024-0512",
     "COUR-2024-007",
     "Nous avons bien pris connaissance de votre appel d'offres n°AO-2024-0512 publié le "
     "10 octobre 2024, dont la date limite de remise des offres est fixée au 25 octobre 2024.",
     "La complexité du cahier des charges et la nécessité de consulter plusieurs partenaires "
     "techniques ne nous permettent pas de préparer une offre de qualité dans ce délai. "
     "Nous sollicitons donc un report de la date de remise au 8 novembre 2024, "
     "délai qui nous paraît raisonnable pour vous soumettre une proposition véritablement compétitive.",
     "Nous espérons que vous accueillerez favorablement cette demande et vous en remercions par avance."),
    ("Convocation – Réunion du Comité de Direction du 15 novembre 2024",
     "COUR-2024-008",
     "Vous êtes convoqué(e) à la prochaine réunion du Comité de Direction qui se tiendra le "
     "vendredi 15 novembre 2024 à 9h30 dans la Salle Panoramique (4ème étage, bâtiment A).",
     "L'ordre du jour est le suivant :\n"
     "1. Approbation du compte-rendu du 18 octobre 2024\n"
     "2. Point sur les résultats financiers T3 2024\n"
     "3. Présentation du plan stratégique 2025-2027\n"
     "4. Ressources humaines : plan de recrutement Q1 2025\n"
     "5. Questions diverses\n\n"
     "Merci de confirmer votre présence avant le 12 novembre 2024.",
     "Les documents de travail vous seront transmis au moins 48h avant la réunion."),
    ("Accusé de réception – Dossier de candidature n°CAND-2024-0089",
     "COUR-2024-009",
     "Nous avons bien reçu votre dossier de candidature n°CAND-2024-0089, "
     "transmis par courrier recommandé le 18 octobre 2024.",
     "Votre dossier est en cours d'instruction par nos services. Nous vous informerons "
     "de la suite donnée à votre candidature dans un délai de 4 semaines. "
     "Toute pièce manquante identifiée lors de l'examen de recevabilité vous sera signalée "
     "par retour de courrier.",
     "Nous vous remercions de l'intérêt que vous portez à notre organisation."),
    ("Lettre de recommandation – Mme Amandine Leroy",
     "COUR-2024-010",
     "J'ai l'honneur de vous recommander chaleureusement Mme Amandine Leroy, "
     "qui a exercé les fonctions de Directrice de Projet au sein de notre organisation "
     "du 1er janvier 2021 au 30 septembre 2024.",
     "Durant ces 3 années et 9 mois, Mme Leroy a démontré des compétences exceptionnelles "
     "en gestion de projets complexes, management d'équipe et relation client. "
     "Elle a notamment piloté avec succès la migration de notre système d'information, "
     "un projet de 1,2 M€ livré dans les délais et sous budget.",
     "C'est donc avec la plus grande confiance que je vous recommande Mme Leroy pour tout poste "
     "à responsabilité en gestion de projet ou direction de programme."),
]

def gen_courrier(idx, sujet, ref, intro, corps, conclusion):
    expediteur_org = ENTREPRISES[idx % len(ENTREPRISES)]
    expediteur = PERSONNES[idx % len(PERSONNES)]
    destinataire_org = ENTREPRISES[(idx + 6) % len(ENTREPRISES)]
    destinataire = PERSONNES[(idx + 3) % len(PERSONNES)]

    doc = Document()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(
        f"{expediteur['nom']}\n{expediteur['poste']}\n{expediteur_org['nom']}\n"
        f"{expediteur_org['adresse']}\n{expediteur_org['tel']}\n{expediteur_org['email']}"
    )

    doc.add_paragraph()
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.add_run(f"{expediteur_org['ville']}, le {date_str()}")

    doc.add_paragraph()
    doc.add_paragraph("À l'attention de :")
    doc.add_paragraph(
        f"{destinataire['nom']}\n{destinataire['poste']}\n"
        f"{destinataire_org['nom']}\n{destinataire_org['adresse']}"
    )
    doc.add_paragraph()
    doc.add_paragraph(f"Réf. : {ref}")
    doc.add_paragraph(f"Objet : {sujet}")
    doc.add_paragraph()
    doc.add_paragraph(f"Madame, Monsieur,")
    doc.add_paragraph()
    doc.add_paragraph(intro)
    doc.add_paragraph()
    doc.add_paragraph(corps)
    doc.add_paragraph()
    doc.add_paragraph(conclusion)
    doc.add_paragraph()
    doc.add_paragraph(
        "Dans l'attente de votre retour, nous vous prions d'agréer, Madame, Monsieur, "
        "l'expression de nos salutations distinguées."
    )
    doc.add_paragraph()
    doc.add_paragraph(f"{expediteur['nom']}\n{expediteur['poste']}\n{expediteur_org['nom']}")

    path = os.path.join(BASE_DIR, "courriers_administratifs", f"courrier_{idx:03d}.docx")
    doc.save(path)
    return path

# ────────────────────────────────────────────────
# PROCÈS-VERBAUX (10 docs)
# ────────────────────────────────────────────────

PV_DATA = [
    ("PV-CA-2024-001", "Conseil d'Administration", "15/01/2024", "09h30", "12h15",
     "Salle Panoramique – Siège social",
     [("Approbation des comptes 2023", "Les comptes de l'exercice 2023 présentant un résultat net de 850 000 € sont approuvés à l'unanimité.", "APPROUVÉ à l'unanimité."),
      ("Affectation du résultat 2023", "Après délibération, il est décidé d'affecter 200 000 € en réserves légales et 650 000 € en report à nouveau.", "DÉCIDÉ à l'unanimité.")]),
    ("PV-CODIR-2024-001", "Comité de Direction", "08/02/2024", "10h00", "12h30",
     "Salle Einstein – 4ème étage",
     [("Résultats T4 2023 et plan 2024", "Présentation des résultats T4 par le DAF. CA de 3,2 M€ en hausse de 8 % vs T4 2022. Objectif 2024 fixé à 14 M€.", "OBJECTIF 2024 validé à 14 M€."),
      ("Recrutements prioritaires Q1 2024", "La DRH propose 4 recrutements : 2 développeurs seniors, 1 chef de projet, 1 analyste BI.", "BUDGET RECRUTEMENT approuvé : 180 000 € annuels chargés.")]),
    ("PV-AGO-2024-001", "Assemblée Générale Ordinaire", "28/03/2024", "14h00", "17h00",
     "Grande Salle de Conférence – 1er étage",
     [("Approbation des comptes sociaux 2023", "Les comptes sociaux présentés par le Président sont adoptés (CA HT : 12,45 M€, résultat net : 850 K€).", "ADOPTÉ à l'unanimité."),
      ("Renouvellement du mandat du Commissaire aux Comptes", "Le mandat du cabinet Deloitte & Associés est renouvelé pour une durée de 6 exercices.", "RENOUVELÉ à l'unanimité.")]),
    ("PV-CSE-2024-001", "Réunion du CSE – Consultation trimestrielle", "12/04/2024", "09h00", "11h30",
     "Salle de Réunion RH",
     [("Point sur la situation économique et financière", "Le DAF présente le bilan T1 2024 : CA 3,45 M€ (+7,8 %), EBIT 12,1 %. Effectif : 87 CDI, 3 CDD, 4 alternants.", "INFORMATION actée."),
      ("Projet de réorganisation du service clients", "La Direction présente le projet de fusion des équipes support et relation client. 3 postes impactés avec reclassement interne.", "AVIS FAVORABLE sous réserve d'accompagnement RH individualisé.")]),
    ("PV-PROJ-2024-001", "Réunion de projet – Projet OMEGA", "22/05/2024", "14h00", "16h00",
     "Visioconférence Teams",
     [("Avancement du lot 2 – Intégration données", "L'équipe technique présente l'état d'avancement : 68 % des travaux réalisés, retard de 1 semaine sur le planning initial.", "PLAN DE RATTRAPAGE validé – livraison lot 2 repoussée au 14 juin 2024."),
      ("Budget – Point de suivi", "Budget consommé : 142 000 € sur 180 000 € alloués. Reste à engager : 38 000 €. Prévision de dépassement : 12 000 €.", "DEMANDE D'AVENANT approuvée pour 12 000 € supplémentaires.")]),
    ("PV-CRISE-2024-001", "Cellule de crise – Incident informatique P1", "03/06/2024", "08h30", "13h00",
     "War Room – Salle Informatique B2",
     [("Analyse de l'incident – Panne datacenter", "Panne du datacenter principal à 06h47. Bascule automatique sur site de secours à 07h12. Durée d'interruption : 25 minutes. Cause : défaillance UPS.", "PLAN D'ACTION déclenché : audit UPS, test mensuel PRA, rapport post-incident sous 5 jours."),
      ("Communication de crise", "Envoi d'un email d'information aux clients impactés (347 comptes). Hot-line dédiée ouverte jusqu'au 05/06/2024.", "COMMUNIQUÉ DE CRISE approuvé et diffusé à 10h00.")]),
    ("PV-CA-2024-002", "Conseil d'Administration – Session extraordinaire", "18/07/2024", "15h00", "18h30",
     "Salle Panoramique – Siège social",
     [("Projet d'acquisition de la société Novéa Services", "Présentation du projet d'acquisition par le DG. Valorisation cible : 2,8 M€. Synergies attendues : 400 K€/an à horizon 3 ans.", "PROCESSUS D'ACQUISITION autorisé. Budget due diligence : 85 000 €."),
      ("Modification des délégations de pouvoirs", "Révision des seuils d'engagement financier : DG 500 K€ (vs 200 K€), DAF 200 K€ (vs 100 K€), Directeurs 50 K€.", "NOUVEAU TABLEAU DES DÉLÉGATIONS approuvé.")]),
    ("PV-CODIR-2024-002", "Comité de Direction – Revue stratégique S1", "30/08/2024", "09h00", "13h00",
     "Salle Einstein – 4ème étage",
     [("Résultats S1 2024 vs budget", "CA S1 : 6,65 M€ (+4,2 % vs S1 2023). EBIT S1 : 9,8 % (vs objectif 10,5 %). Retard principalement dû au report de 3 contrats importants.", "RÉVISION DES OBJECTIFS S2 : maintien CA annuel 14 M€, EBIT cible ajusté à 9,5 %."),
      ("Lancement initiative RSE 2024-2026", "Présentation du plan RSE : réduction de l'empreinte carbone de 30 % d'ici 2026, labellisation ISO 14001.", "PLAN RSE approuvé. Budget alloué : 120 000 € sur 3 ans.")]),
    ("PV-REMISE-2024-001", "PV de remise de clés – Bail commercial BAIL-COM-2024-001", "01/09/2024", "10h00", "11h30",
     "8 rue du Temple, 75004 Paris",
     [("Remise des clés et de l'accès", "Le bailleur remet au preneur : 3 clés entrée principale, 2 badges d'accès immeuble, 1 télécommande volet roulant, codes alarme.", "REMISE EFFECTUÉE. Preneur confirme avoir pris possession des locaux."),
      ("État des lieux d'entrée", "Relevé des compteurs : électricité 43 821 kWh, eau 2 156 m³. Aucune réserve émise par les deux parties sur l'état général des locaux.", "ÉTAT DES LIEUX APPROUVÉ par les deux parties.")]),
    ("PV-RECEP-2024-001", "PV de réception de travaux – Rénovation bureaux R+3", "20/10/2024", "09h00", "12h00",
     "Bâtiment principal – 3ème étage",
     [("Vérification des travaux réalisés", "Les travaux de rénovation (peinture, faux-plafonds, réseau électrique) ont été inspectés en présence du maître d'œuvre. 2 réserves mineures identifiées : reprises de peinture couloir nord, remplacement d'un spot défectueux.", "RÉCEPTION PRONONCÉE avec réserves. Levée des réserves avant le 30/10/2024."),
      ("Validation des garanties", "L'entreprise titulaire remet les attestations d'assurance décennale, biennale et dommages-ouvrage.", "GARANTIES VALIDÉES. Dossier archivé.")]),
]

def gen_pv(idx, ref, titre, date_r, h_debut, h_fin, lieu, points):
    org = ENTREPRISES[idx % len(ENTREPRISES)]
    president = PERSONNES[idx % len(PERSONNES)]
    secretaire = PERSONNES[(idx + 2) % len(PERSONNES)]
    participants = [PERSONNES[(idx + i) % len(PERSONNES)] for i in range(5)]

    doc = Document()
    add_header(doc, org["nom"], titre.upper(), ref)

    doc.add_paragraph(
        f"Date : {date_r}\nHeure de début : {h_debut}\nHeure de fin : {h_fin}\n"
        f"Lieu : {lieu}\nPrésidé par : {president['nom']}, {president['poste']}"
    )

    doc.add_heading("PARTICIPANTS", level=2)
    doc.add_paragraph("Présents :")
    for p in participants[:4]:
        doc.add_paragraph(f"  – {p['nom']}, {p['poste']}", style="List Bullet")
    doc.add_paragraph(f"Excusé(e) : {participants[4]['nom']}, {participants[4]['poste']}")
    doc.add_paragraph(f"Secrétaire de séance : {secretaire['nom']}, {secretaire['poste']}")

    doc.add_heading("ORDRE DU JOUR", level=2)
    doc.add_paragraph("1. Approbation du procès-verbal de la séance précédente")
    for i, (sujet, *_) in enumerate(points, 2):
        doc.add_paragraph(f"{i}. {sujet}")
    doc.add_paragraph(f"{len(points) + 2}. Questions diverses")

    doc.add_heading("DÉROULEMENT DE LA SÉANCE", level=2)

    doc.add_heading("1. Approbation du PV précédent", level=3)
    doc.add_paragraph(
        f"Le procès-verbal de la séance précédente est soumis à approbation. "
        f"Adopté à l'unanimité des membres présents."
    )

    for i, (sujet, cr, decision) in enumerate(points, 2):
        doc.add_heading(f"{i}. {sujet}", level=3)
        doc.add_paragraph(cr)
        p = doc.add_paragraph(f"→ Décision : {decision}")
        p.runs[0].bold = True

    doc.add_heading(f"{len(points) + 2}. Questions diverses", level=3)
    doc.add_paragraph("Aucune question diverse n'a été soulevée.")

    doc.add_heading("CLÔTURE DE LA SÉANCE", level=2)
    doc.add_paragraph(f"La séance est levée à {h_fin} par {president['nom']}.")

    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = f"Le Président de séance\n{president['nom']}\nSignature : _______________"
    table.rows[0].cells[1].text = f"Le Secrétaire de séance\n{secretaire['nom']}\nSignature : _______________"

    path = os.path.join(BASE_DIR, "proces_verbaux", f"pv_{idx:03d}.docx")
    doc.save(path)
    return path

# ────────────────────────────────────────────────
# BONS DE COMMANDE (10 docs)
# ────────────────────────────────────────────────

BDC_DATA = [
    ("BDC-2024-0001", "Matériel informatique – Renouvellement 10 postes HP EliteBook", 15600, 10, "Renouvellement parc informatique vieillissant"),
    ("BDC-2024-0002", "Mobilier de bureau – Tables et chaises Salle de Réunion A", 4200, 1, "Aménagement salle réunion restructurée"),
    ("BDC-2024-0003", "Consommables imprimante – Cartouches Canon PG-540 × 200", 890, 200, "Stock annuel consommables reprographique"),
    ("BDC-2024-0004", "Licences Microsoft 365 Business Premium – 20 postes", 3600, 20, "Renouvellement annuel licences suite collaborative"),
    ("BDC-2024-0005", "Hébergement cloud Azure – Abonnement annuel 2025", 7200, 1, "Migration cloud et hébergement applicatif"),
    ("BDC-2024-0006", "Switch réseau Cisco Catalyst 2960 + câblage Cat.6", 5400, 3, "Modernisation infrastructure réseau LAN"),
    ("BDC-2024-0007", "Formation Cybersécurité ISO 27001 – 5 collaborateurs", 2800, 5, "Montée en compétences équipe SI"),
    ("BDC-2024-0008", "Papeterie et fournitures – Commande trimestrielle Q4", 650, 1, "Approvisionnement trimestriel fournitures bureautiques"),
    ("BDC-2024-0009", "Système vidéosurveillance IP 8 caméras + NVR", 3900, 1, "Mise en conformité sécurité physique bâtiment"),
    ("BDC-2024-0010", "Serveur NAS Synology DS1823xs+ + disques 8 To × 8", 12000, 1, "Infrastructure sauvegarde et stockage données"),
]

def gen_bdc(idx, num, objet, montant_ht, qte, motif):
    acheteur = ENTREPRISES[idx % len(ENTREPRISES)]
    fournisseur = ENTREPRISES[(idx + 7) % len(ENTREPRISES)]
    responsable = PERSONNES[idx % len(PERSONNES)]
    approbateur = PERSONNES[(idx + 1) % len(PERSONNES)]
    tva = round(montant_ht * 0.20, 2)
    ttc = round(montant_ht + tva, 2)
    pu = round(montant_ht / qte, 2) if qte > 0 else montant_ht

    doc = Document()
    add_header(doc, acheteur["nom"], "BON DE COMMANDE", num)

    doc.add_paragraph(
        f"Acheteur : {acheteur['nom']}\nSIRET : {acheteur['siret']}\n"
        f"Responsable achat : {responsable['nom']}, {responsable['poste']}\n"
        f"Date de commande : {date_str()}\n"
        f"Date de livraison souhaitée : {date_str(14)}\n"
        f"Motif : {motif}"
    )
    doc.add_paragraph()
    doc.add_paragraph(
        f"Fournisseur : {fournisseur['nom']}\nSIRET : {fournisseur['siret']}\n"
        f"Contact : {PERSONNES[(idx + 5) % len(PERSONNES)]['nom']}\n"
        f"Tél : {fournisseur['tel']}\nEmail : {fournisseur['email']}"
    )
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    for i, h in enumerate(["Réf.", "Désignation", "Qté", "Prix unit. HT", "TVA", "Total HT"]):
        table.rows[0].cells[i].text = h
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    row = table.add_row().cells
    row[0].text = f"REF-{num[-4:]}"
    row[1].text = objet
    row[2].text = str(qte)
    row[3].text = f"{pu:,.2f} €"
    row[4].text = "20 %"
    row[5].text = f"{montant_ht:,.2f} €"

    doc.add_paragraph()
    doc.add_paragraph(f"Total HT : {montant_ht:,.2f} €")
    doc.add_paragraph(f"TVA (20 %) : {tva:,.2f} €")
    p = doc.add_paragraph(f"TOTAL TTC : {ttc:,.2f} €")
    p.runs[0].bold = True

    doc.add_paragraph()
    doc.add_paragraph(
        f"Adresse de livraison : {acheteur['adresse']}\n"
        f"Instructions : Livraison en semaine, horaires 8h-18h, contacter {responsable['nom']} – {responsable['tel']}\n"
        f"Conditions de paiement : 30 jours net date de facture, virement {acheteur['banque']}"
    )
    doc.add_paragraph()
    doc.add_paragraph(
        f"Bon de commande approuvé par : {approbateur['nom']}\n"
        f"Fonction : {approbateur['poste']}\n"
        f"Date d'approbation : {date_str()}\n"
        f"Signature : _______________"
    )

    path = os.path.join(BASE_DIR, "bons_de_commande", f"bdc_{idx:03d}.docx")
    doc.save(path)
    return path

# ────────────────────────────────────────────────
# BULLETINS DE SALAIRE (10 docs)
# ────────────────────────────────────────────────

BULLETINS_DATA = [
    ("BULL-2024-JAN-001", "janvier 2024", 4500, "Développeur Senior", "75", 218, "DEV-2024-042"),
    ("BULL-2024-FEV-001", "février 2024", 4200, "Chef de projet", "75", 218, "CP-2024-017"),
    ("BULL-2024-MAR-001", "mars 2024", 3800, "Analyste financier", "75", 218, "AF-2024-031"),
    ("BULL-2024-AVR-001", "avril 2024", 3600, "Responsable RH", "75", 218, "RRH-2024-008"),
    ("BULL-2024-MAI-001", "mai 2024", 5500, "Directeur Commercial", "75", 218, "DC-2024-002"),
    ("BULL-2024-JUN-001", "juin 2024", 3200, "Comptable", "75", 218, "COMPT-2024-019"),
    ("BULL-2024-JUL-001", "juillet 2024", 4100, "Ingénieur Systèmes", "69", 218, "IS-2024-025"),
    ("BULL-2024-AOU-001", "août 2024", 2900, "Chargé(e) de communication", "13", 217, "COMM-2024-013"),
    ("BULL-2024-SEP-001", "septembre 2024", 2800, "Office Manager", "31", 217, "OM-2024-007"),
    ("BULL-2024-OCT-001", "octobre 2024", 4000, "Juriste d'entreprise", "67", 218, "JUR-2024-011"),
]

EMPLOYES = [
    ("Lucas Martin", "12 rue des Acacias", "75017 Paris", "FR76 3000 4005 0123 4567 8901 234"),
    ("Marie Dubois", "8 avenue Foch", "75116 Paris", "FR76 1027 8061 0000 0301 2456 789"),
    ("Julien Bernard", "45 rue Voltaire", "69003 Lyon", "FR76 1780 6004 0300 1234 0056 789"),
    ("Camille Thomas", "3 allée des Pins", "33600 Pessac", "FR76 3000 3032 0500 0206 1345 012"),
    ("Nicolas Robert", "22 boulevard Victor Hugo", "92200 Neuilly-sur-Seine", "FR76 1025 7020 0100 2345 6789 012"),
    ("Aurélie Petit", "18 rue de la Paix", "13001 Marseille", "FR76 1558 9200 0600 3456 7890 123"),
    ("Thomas Durand", "7 impasse du Moulin", "31000 Toulouse", "FR76 1660 6001 0200 4567 8901 234"),
    ("Emma Lefevre", "15 rue Nationale", "59000 Lille", "FR76 3000 4000 0700 5678 9012 345"),
    ("Hugo Morel", "9 place du Marché", "67000 Strasbourg", "FR76 1466 7000 0300 6789 0123 456"),
    ("Léa Simon", "34 cours Pasteur", "33000 Bordeaux", "FR76 3000 3032 0800 7890 1234 567"),
]

def gen_bulletin(idx, num, periode, brut, poste, dept, nb_heures, matricule):
    org = ENTREPRISES[idx % len(ENTREPRISES)]
    employe = EMPLOYES[idx % len(EMPLOYES)]
    nom_emp, adresse_emp, cp_ville_emp, iban_emp = employe

    # Calculs
    ss_sal = round(brut * 0.069, 2)
    ret_comp = round(brut * 0.0315, 2)
    chomage = round(brut * 0.024, 2)
    mutuelle_sal = round(brut * 0.015, 2)
    csg_ded = round(brut * 0.068, 2)
    csg_non_ded = round(brut * 0.029, 2)
    total_cotis_sal = round(ss_sal + ret_comp + chomage + mutuelle_sal + csg_ded + csg_non_ded, 2)
    net_imposable = round(brut - total_cotis_sal + csg_non_ded, 2)
    net_a_payer = round(brut - total_cotis_sal, 2)

    cotis_pat_ss = round(brut * 0.128, 2)
    cotis_pat_ret = round(brut * 0.059, 2)
    cotis_pat_chomage = round(brut * 0.042, 2)
    cotis_pat_accident = round(brut * 0.022, 2)
    cotis_pat_mutuelle = round(brut * 0.015, 2)
    total_cotis_pat = round(cotis_pat_ss + cotis_pat_ret + cotis_pat_chomage + cotis_pat_accident + cotis_pat_mutuelle, 2)
    cout_total = round(brut + total_cotis_pat, 2)

    doc = Document()
    add_header(doc, org["nom"], "BULLETIN DE PAIE", num)

    # Bloc employeur / employé
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    t.rows[0].cells[0].text = (
        f"EMPLOYEUR :\n{org['nom']}\n{org['adresse']}\n"
        f"SIRET : {org['siret']}\nNAF : 6201Z\n"
        f"Convention collective : SYNTEC (IDCC 1486)"
    )
    t.rows[0].cells[1].text = (
        f"EMPLOYÉ(E) :\n{nom_emp}\n{adresse_emp}\n{cp_ville_emp}\n"
        f"Matricule : {matricule}\nPoste : {poste}\n"
        f"Département : {dept}\nDate d'entrée : 01/03/2022"
    )

    doc.add_paragraph()
    doc.add_paragraph(
        f"Période : {periode}  |  Nb heures : {nb_heures}h  |  "
        f"Date de paiement : {date_str(-2)}"
    )
    doc.add_paragraph()

    # Tableau des éléments de paie
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    for i, h in enumerate(["Rubrique", "Base", "Taux salarial", "Taux patronal", "Montant (€)"]):
        table.rows[0].cells[i].text = h
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    lignes = [
        ("Salaire de base", f"{nb_heures}h", "-", "-", f"{brut:,.2f}"),
        ("Cotis. Sécu. sociale (maladie)", f"{brut:,.2f} €", "6,90 %", "12,80 %", f"-{ss_sal:,.2f}"),
        ("Retraite complémentaire AGIRC-ARRCO", f"{brut:,.2f} €", "3,15 %", "5,90 %", f"-{ret_comp:,.2f}"),
        ("Assurance chômage", f"{brut:,.2f} €", "2,40 %", "4,20 %", f"-{chomage:,.2f}"),
        ("Mutuelle santé complémentaire", f"{brut:,.2f} €", "1,50 %", "1,50 %", f"-{mutuelle_sal:,.2f}"),
        ("CSG déductible", f"{brut:,.2f} €", "6,80 %", "-", f"-{csg_ded:,.2f}"),
        ("CSG/CRDS non déductible", f"{brut:,.2f} €", "2,90 %", "-", f"-{csg_non_ded:,.2f}"),
        ("Cotis. AT/MP", f"{brut:,.2f} €", "-", "2,20 %", "0,00"),
    ]
    for l in lignes:
        r = table.add_row().cells
        for i, v in enumerate(l):
            r[i].text = v

    doc.add_paragraph()
    doc.add_paragraph(f"Salaire brut total : {brut:,.2f} €")
    doc.add_paragraph(f"Total cotisations salariales : -{total_cotis_sal:,.2f} €")
    doc.add_paragraph(f"Net imposable : {net_imposable:,.2f} €")
    p = doc.add_paragraph(f"NET À PAYER AVANT IMPÔT : {net_a_payer:,.2f} €")
    p.runs[0].bold = True
    doc.add_paragraph(f"Prélèvement à la source (taux personnalisé) : -{round(net_imposable * 0.08, 2):,.2f} €")
    doc.add_paragraph(f"NET À PAYER APRÈS PRÉLÈVEMENT : {round(net_a_payer - net_imposable * 0.08, 2):,.2f} €")
    doc.add_paragraph(f"\nCoût total employeur : {cout_total:,.2f} €")
    doc.add_paragraph(
        f"\nVirement effectué le {date_str(-2)} – "
        f"IBAN : {iban_emp}"
    )

    path = os.path.join(BASE_DIR, "bulletins_de_salaire", f"bulletin_{idx:03d}.docx")
    doc.save(path)
    return path

# ────────────────────────────────────────────────
# RELEVÉS BANCAIRES (10 docs)
# ────────────────────────────────────────────────

RELEVES_DATA = [
    ("REL-BNQ-2024-JAN", "janvier 2024", "01/01/2024", "31/01/2024", 145800, 182450),
    ("REL-BNQ-2024-FEV", "février 2024", "01/02/2024", "29/02/2024", 182450, 167300),
    ("REL-BNQ-2024-MAR", "mars 2024", "01/03/2024", "31/03/2024", 167300, 221600),
    ("REL-BNQ-2024-AVR", "avril 2024", "01/04/2024", "30/04/2024", 221600, 198700),
    ("REL-BNQ-2024-MAI", "mai 2024", "01/05/2024", "31/05/2024", 198700, 245100),
    ("REL-BNQ-2024-JUN", "juin 2024", "01/06/2024", "30/06/2024", 245100, 213400),
    ("REL-BNQ-2024-JUL", "juillet 2024", "01/07/2024", "31/07/2024", 213400, 189900),
    ("REL-BNQ-2024-AOU", "août 2024", "01/08/2024", "31/08/2024", 189900, 204300),
    ("REL-BNQ-2024-SEP", "septembre 2024", "01/09/2024", "30/09/2024", 204300, 268700),
    ("REL-BNQ-2024-OCT", "octobre 2024", "01/10/2024", "31/10/2024", 268700, 235500),
]

def gen_releve(idx, num, periode, date_debut, date_fin, solde_ouv, solde_clo):
    org = ENTREPRISES[idx % len(ENTREPRISES)]
    banque = ["BNP Paribas", "Société Générale", "Crédit Agricole", "LCL", "CIC"][idx % 5]
    agence = ["Paris 11ème", "Lyon Part-Dieu", "Marseille Canebière", "Bordeaux Centre", "Lille Grand Place"][idx % 5]

    clients = [ENTREPRISES[(idx + i) % len(ENTREPRISES)]["nom"] for i in range(3)]
    fournisseurs = [ENTREPRISES[(idx + i + 5) % len(ENTREPRISES)]["nom"] for i in range(3)]

    ops = [
        (f"02/{date_debut[3:]}", f"VIR ENTRANT – {clients[0]}", None, round((solde_clo - solde_ouv) * 0.35 + 45000)),
        (f"05/{date_debut[3:]}", f"PRELEVEMENT – LOYER BUREAU {org['ville'].upper()}", round(org['capital'].replace(' ', '').replace('€', '').__len__() * 200 + 1400), None),
        (f"08/{date_debut[3:]}", f"VIR SORTANT – {fournisseurs[0]}", round((solde_clo - solde_ouv + 80000) * 0.20), None),
        (f"10/{date_debut[3:]}", f"VIR ENTRANT – {clients[1]}", None, round((solde_clo - solde_ouv) * 0.25 + 28000)),
        (f"12/{date_debut[3:]}", "PRELEVEMENT – URSSAF COTISATIONS", round(solde_ouv * 0.08), None),
        (f"15/{date_debut[3:]}", "PAIEMENT CB – FOURNITURES ET SERVICES", round(800 + idx * 50), None),
        (f"18/{date_debut[3:]}", f"VIR SORTANT – SALAIRES {periode.upper()}", round(solde_ouv * 0.12), None),
        (f"20/{date_debut[3:]}", "FRAIS BANCAIRES MENSUELS", round(180 + idx * 10), None),
        (f"22/{date_debut[3:]}", "VIR ENTRANT – REMBOURSEMENT TVA DGFiP", None, round(solde_ouv * 0.03)),
        (f"25/{date_debut[3:]}", f"VIR ENTRANT – {clients[2]}", None, round((solde_clo - solde_ouv) * 0.40 + 35000)),
        (f"28/{date_debut[3:]}", f"VIR SORTANT – {fournisseurs[1]}", round(8500 + idx * 300), None),
        (f"30/{date_debut[3:]}", f"VIR SORTANT – {fournisseurs[2]}", round(3200 + idx * 150), None),
    ]

    total_debits = sum(o[2] for o in ops if o[2]) or 1
    total_credits = sum(o[3] for o in ops if o[3]) or 1

    doc = Document()
    add_header(doc, banque, f"RELEVÉ DE COMPTE – {periode.upper()}", num)

    doc.add_paragraph(
        f"Titulaire : {org['nom']}\nAdresse : {org['adresse']}\n"
        f"N° de compte : {org['iban'][4:25].replace(' ', '')[:11]}\n"
        f"IBAN : {org['iban']}\nBIC : {org['bic']}\n"
        f"Agence : {banque} – {agence}"
    )
    doc.add_paragraph()
    p = doc.add_paragraph(f"Solde au {date_debut} : {solde_ouv:,.2f} €")
    p.runs[0].bold = True
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for i, h in enumerate(["Date", "Libellé opération", "Débit (€)", "Crédit (€)"]):
        table.rows[0].cells[i].text = h
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    for date_op, libelle, debit, credit in ops:
        r = table.add_row().cells
        r[0].text = date_op
        r[1].text = libelle
        r[2].text = f"{debit:,.2f}" if debit else ""
        r[3].text = f"{credit:,.2f}" if credit else ""

    doc.add_paragraph()
    doc.add_paragraph(f"Total débits : {total_debits:,.2f} €  |  Total crédits : {total_credits:,.2f} €")
    p2 = doc.add_paragraph(f"Solde au {date_fin} : {solde_clo:,.2f} €")
    p2.runs[0].bold = True
    doc.add_paragraph()
    doc.add_paragraph(
        "Ce document est émis à titre informatif. En cas d'anomalie, veuillez contacter "
        f"votre agence {agence} sous 30 jours."
    )

    path = os.path.join(BASE_DIR, "releves_bancaires", f"releve_{idx:03d}.docx")
    doc.save(path)
    return path

# ────────────────────────────────────────────────
# CONVENTIONS (10 docs)
# ────────────────────────────────────────────────

CONVENTIONS_DATA = [
    ("CONV-PART-2024-001", "Convention de partenariat stratégique",
     "Les deux parties souhaitent formaliser leur coopération dans le cadre du développement commercial en Europe du Sud.",
     "Développer conjointement une offre de services intégrée à destination des PME industrielles.",
     "Partie A : prospection commerciale et avant-vente (minimum 5 rendez-vous qualifiés/mois).",
     "Partie B : réponse technique, démonstrations, déploiement et support clients.",
     "3 ans", "30 jours", "Partage des revenus : 60 % Partie A, 40 % Partie B sur les contrats co-signés."),
    ("CONV-STAGE-2024-001", "Convention de stage",
     "Dans le cadre de la formation de l'étudiant, l'entreprise accueille celui-ci pour une période de stage.",
     "Permettre à l'étudiant d'acquérir une expérience professionnelle dans le domaine de l'analyse de données.",
     "Partie A (entreprise) : accueil, encadrement, accès aux outils, attribution d'un tuteur désigné.",
     "Partie B (établissement) : suivi pédagogique, validation des compétences, délivrance du diplôme.",
     "4 mois (du 03/02/2025 au 30/05/2025)", "8 jours", "Gratification légale : 4,35 €/heure (soit ~676 €/mois)."),
    ("CONV-MAD-2024-001", "Convention de mise à disposition de personnel",
     "La Partie A consent à mettre à disposition de la Partie B un de ses salariés pour une mission définie.",
     "Renforcer temporairement les équipes de la Partie B lors d'un pic d'activité identifié.",
     "Partie A : maintien du lien contractuel avec le salarié, paiement de la rémunération.",
     "Partie B : direction fonctionnelle, fourniture des outils de travail, refacturation des coûts salariaux à la Partie A.",
     "6 mois renouvelable 1 fois", "15 jours",
     "Refacturation mensuelle du coût salarial chargé majoré de 5 % de frais de gestion."),
    ("CONV-MECENAT-2024-001", "Convention de mécénat",
     "Dans le cadre de sa politique RSE, la Partie A souhaite soutenir les actions menées par la Partie B.",
     "Financer le programme d'insertion professionnelle de jeunes en difficulté mené par la Partie B.",
     "Partie A : versement d'un don annuel de 25 000 €, aide en nature (équipements informatiques).",
     "Partie B : utilisation des fonds conformément à l'objet, rapport annuel d'activité, mise en avant du mécène.",
     "3 ans (2024-2026)", "30 jours",
     "Don annuel de 25 000 € déductible fiscalement à hauteur de 60 % (art. 238 bis CGI)."),
    ("CONV-CADRE-2024-001", "Convention cadre de collaboration",
     "Les parties souhaitent établir un cadre général régissant l'ensemble de leurs relations contractuelles.",
     "Définir les principes et conditions générales applicables à l'ensemble des commandes et prestations entre les parties.",
     "Partie A : donneur d'ordre, validation des commandes, paiement dans les délais convenus.",
     "Partie B : exécution des prestations selon les spécifications, respect des délais, reporting régulier.",
     "2 ans avec reconduction tacite", "60 jours",
     "Conditions tarifaires négociées annuellement, remise de volume de 5 % au-delà de 100 K€/an."),
    ("CONV-CESSION-2024-001", "Convention de cession de droits de propriété intellectuelle",
     "La Partie A est titulaire de droits de propriété intellectuelle qu'elle souhaite céder à la Partie B.",
     "Céder l'ensemble des droits sur le logiciel « AtlasDoc v2.1 » à la Partie B pour exploitation commerciale.",
     "Partie A (cédant) : garantie d'éviction, documentation technique complète, formation initiale.",
     "Partie B (cessionnaire) : paiement du prix de cession, non-divulgation du code source.",
     "Cession définitive et irrévocable", "N/A",
     "Prix de cession : 180 000 € HT, payable en 3 tranches de 60 000 €."),
    ("CONV-PRET-2024-001", "Convention de prêt de matériel",
     "La Partie A met à la disposition de la Partie B du matériel informatique pour une durée déterminée.",
     "Permettre à la Partie B d'utiliser 5 ordinateurs portables et 2 serveurs lors de son déménagement.",
     "Partie A : livraison du matériel en bon état de fonctionnement, assistance technique initiale.",
     "Partie B : utilisation conforme, entretien courant, restitution dans l'état initial, assurance.",
     "3 mois", "8 jours",
     "Prêt à titre gratuit. La Partie B prend en charge l'assurance du matériel pour la durée du prêt."),
    ("CONV-SUBV-2024-001", "Convention de subvention",
     "L'autorité financière (Partie A) accorde une subvention à l'organisme bénéficiaire (Partie B).",
     "Soutenir la réalisation du projet de modernisation du système de gestion documentaire.",
     "Partie A : versement de la subvention selon l'échéancier, contrôle de l'utilisation des fonds.",
     "Partie B : réalisation du projet selon le cahier des charges, production de justificatifs, rapport final.",
     "18 mois", "N/A",
     "Subvention de 120 000 € versée en 3 tranches : 40 % à la signature, 40 % à mi-parcours, 20 % à la clôture."),
    ("CONV-ALTERN-2024-001", "Convention de formation en alternance (apprentissage)",
     "L'entreprise accueille un apprenti dans le cadre d'un contrat d'apprentissage.",
     "Former un alternant au métier de développeur full-stack en partenariat avec un CFA.",
     "Partie A (entreprise) : maître d'apprentissage, rémunération légale, conditions d'accueil.",
     "Partie B (CFA) : formation théorique, suivi pédagogique, certification professionnelle.",
     "2 ans (septembre 2024 – août 2026)", "30 jours",
     "Rémunération : 65 % du SMIC (1ère année), 80 % du SMIC (2ème année). Aide alternance : 6 000 €."),
    ("CONV-DPA-2024-001", "Convention de traitement de données personnelles (DPA)",
     "Dans le cadre de leur collaboration, les parties échangent des données à caractère personnel.",
     "Encadrer le traitement des données personnelles conformément au RGPD (UE) 2016/679.",
     "Partie A (responsable de traitement) : définition des finalités, base légale, droits des personnes.",
     "Partie B (sous-traitant) : traitement selon les seules instructions de la Partie A, sécurité des données, notification des incidents.",
     "Durée de la prestation principale", "30 jours",
     "Sans contrepartie financière spécifique – intégré au contrat de service principal."),
]

def gen_convention(idx, ref, titre, contexte, objectif, oblig_a, oblig_b, duree, preavis, financier):
    part_a = ENTREPRISES[idx % len(ENTREPRISES)]
    part_b = ENTREPRISES[(idx + 4) % len(ENTREPRISES)]
    rep_a = PERSONNES[idx % len(PERSONNES)]
    rep_b = PERSONNES[(idx + 5) % len(PERSONNES)]
    juridiction = VILLES[idx % len(VILLES)]

    doc = Document()
    add_header(doc, part_a["nom"], titre.upper(), ref)

    doc.add_heading("PARTIES", level=2)
    doc.add_paragraph(
        f"D'une part : {part_a['nom']}, {part_a['forme']} au capital de {part_a['capital']} €, "
        f"SIREN {part_a['siren']}, {part_a['adresse']}, représentée par {rep_a['nom']}, "
        f"{rep_a['poste']}, ci-après dénommée « Partie A »,"
    )
    doc.add_paragraph(
        f"D'autre part : {part_b['nom']}, {part_b['forme']} au capital de {part_b['capital']} €, "
        f"SIREN {part_b['siren']}, {part_b['adresse']}, représentée par {rep_b['nom']}, "
        f"{rep_b['poste']}, ci-après dénommée « Partie B »."
    )

    doc.add_heading("PRÉAMBULE", level=2)
    doc.add_paragraph(contexte)
    doc.add_paragraph(f"Les parties formalisent leur accord afin de : {objectif}")

    doc.add_heading("ARTICLE 1 – OBJET", level=2)
    doc.add_paragraph(f"La présente convention a pour objet de fixer les conditions et modalités de : {titre.lower()}.")

    doc.add_heading("ARTICLE 2 – OBLIGATIONS DES PARTIES", level=2)
    doc.add_paragraph(f"La Partie A ({part_a['nom']}) s'engage à :")
    doc.add_paragraph(oblig_a)
    doc.add_paragraph(f"La Partie B ({part_b['nom']}) s'engage à :")
    doc.add_paragraph(oblig_b)

    doc.add_heading("ARTICLE 3 – DURÉE ET RENOUVELLEMENT", level=2)
    doc.add_paragraph(
        f"La présente convention est conclue pour une durée de {duree}, "
        f"à compter du {date_str()}. Elle est renouvelable par accord express des parties."
    )

    doc.add_heading("ARTICLE 4 – CONDITIONS FINANCIÈRES", level=2)
    doc.add_paragraph(financier)

    doc.add_heading("ARTICLE 5 – CONFIDENTIALITÉ ET DONNÉES PERSONNELLES", level=2)
    doc.add_paragraph(
        "Les parties s'engagent à traiter de manière strictement confidentielle toute information échangée. "
        "Le traitement des données personnelles est régi par le RGPD (UE) 2016/679 et la loi Informatique et Libertés."
    )

    doc.add_heading("ARTICLE 6 – RÉSILIATION", level=2)
    doc.add_paragraph(
        f"Chaque partie peut résilier la présente convention avec un préavis de {preavis} "
        "par lettre recommandée avec accusé de réception. En cas de manquement grave, "
        "la résiliation peut être immédiate après mise en demeure restée sans effet sous 8 jours."
    )

    doc.add_heading("ARTICLE 7 – DROIT APPLICABLE ET JURIDICTION", level=2)
    doc.add_paragraph(
        f"La présente convention est soumise au droit français. "
        f"Tout litige sera porté devant les juridictions compétentes de {juridiction}."
    )

    add_signature_block(doc, part_a["nom"], rep_a["nom"], rep_a["poste"],
                        part_b["nom"], rep_b["nom"], rep_b["poste"], part_a["ville"])
    path = os.path.join(BASE_DIR, "conventions", f"convention_{idx:03d}.docx")
    doc.save(path)
    return path

# ────────────────────────────────────────────────
# MAIN
# ────────────────────────────────────────────────

def main():
    print("Création des répertoires...")
    setup_dirs()

    all_docx = []
    to_convert_pdf = []

    print("Génération des contrats de prestation...")
    for i, (obj, ref, dur, mnt) in enumerate(CONTRAT_PRESTATIONS, 1):
        p = gen_contrat_prestation(i, obj, ref, dur, mnt)
        all_docx.append(p)
        if i <= 3:
            to_convert_pdf.append(p)

    print("Génération des contrats de bail...")
    for i, data in enumerate(BAUX_DATA, 1):
        p = gen_bail(i, *data)
        all_docx.append(p)
        if i <= 3:
            to_convert_pdf.append(p)

    print("Génération des factures...")
    for i, (num, obj, ht) in enumerate(FACTURES_DATA, 1):
        p = gen_facture(i, num, obj, ht)
        all_docx.append(p)
        if i <= 3:
            to_convert_pdf.append(p)

    print("Génération des rapports financiers...")
    for i, data in enumerate(RAPPORTS_FIN_DATA, 1):
        p = gen_rapport_financier(i, *data)
        all_docx.append(p)
        if i <= 3:
            to_convert_pdf.append(p)

    print("Génération des courriers administratifs...")
    for i, (sujet, ref, intro, corps, conclusion) in enumerate(COURRIERS_DATA, 1):
        p = gen_courrier(i, sujet, ref, intro, corps, conclusion)
        all_docx.append(p)
        if i <= 3:
            to_convert_pdf.append(p)

    print("Génération des procès-verbaux...")
    for i, data in enumerate(PV_DATA, 1):
        p = gen_pv(i, *data)
        all_docx.append(p)
        if i <= 3:
            to_convert_pdf.append(p)

    print("Génération des bons de commande...")
    for i, data in enumerate(BDC_DATA, 1):
        p = gen_bdc(i, *data)
        all_docx.append(p)
        if i <= 3:
            to_convert_pdf.append(p)

    print("Génération des bulletins de salaire...")
    for i, data in enumerate(BULLETINS_DATA, 1):
        p = gen_bulletin(i, *data)
        all_docx.append(p)
        if i <= 3:
            to_convert_pdf.append(p)

    print("Génération des relevés bancaires...")
    for i, data in enumerate(RELEVES_DATA, 1):
        p = gen_releve(i, *data)
        all_docx.append(p)
        if i <= 3:
            to_convert_pdf.append(p)

    print("Génération des conventions...")
    for i, data in enumerate(CONVENTIONS_DATA, 1):
        p = gen_convention(i, *data)
        all_docx.append(p)
        if i <= 3:
            to_convert_pdf.append(p)

    print(f"\n{len(all_docx)} documents .docx générés.")

    print(f"\nConversion en PDF de {len(to_convert_pdf)} documents...")
    converted = 0
    for docx_path in to_convert_pdf:
        pdf_path = docx_path.replace(".docx", ".pdf")
        try:
            convert(docx_path, pdf_path)
            converted += 1
            print(f"  ✓ {os.path.basename(pdf_path)}")
        except Exception as e:
            print(f"  ✗ {os.path.basename(docx_path)}: {e}")

    print(f"\n{converted} PDFs générés.")
    print(f"\nTous les documents sont dans : {BASE_DIR}")

if __name__ == "__main__":
    main()
